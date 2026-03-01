"""
One-time script to generate the invoice DOCX template for docxtpl.

Run: python3 create_template.py
Output: templates/invoice-default.docx

After generating, customize the template in Word/Pages — change fonts,
colors, margins, column widths, etc. The Python code only injects data
into the Jinja2 placeholders; all styling lives in the template.
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(SCRIPT_DIR, 'templates', 'invoice-default.docx')


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is not None:
        tblPr.remove(borders)
    borders_elm = tblPr.makeelement(qn('w:tblBorders'), {})
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = borders_elm.makeelement(qn(f'w:{border_name}'), {
            qn('w:val'): 'none', qn('w:sz'): '0',
            qn('w:space'): '0', qn('w:color'): 'auto',
        })
        borders_elm.append(border)
    tblPr.append(borders_elm)


def build_template():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)

    # ===== TITLE =====
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(16)
    run = title.add_run('INVOICE')
    run.bold = True
    run.font.size = Pt(32)

    # ===== HEADER TABLE (2 columns, borderless) =====
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(header_table)

    # Left cell: sender info (RichText placeholder)
    left_cell = header_table.cell(0, 0)
    left_cell.width = Inches(3.4)
    p = left_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('{{r left_header }}')
    run.font.size = Pt(10)

    # Right cell: invoice details + bill to (RichText placeholder)
    right_cell = header_table.cell(0, 1)
    right_cell.width = Inches(3.4)
    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('{{r right_header }}')
    run.font.size = Pt(10)

    # Spacer
    doc.add_paragraph('')

    # ===== ITEMIZED TABLE (built entirely in Python as a Subdoc) =====
    p = doc.add_paragraph()
    run = p.add_run('{{ itemized_table }}')
    run.font.size = Pt(9)

    # ===== FOOTER =====
    doc.add_paragraph('')

    p = doc.add_paragraph()
    run = p.add_run('Make all checks payable to ')
    run.font.size = Pt(10)
    run = p.add_run('{{ payable_to }}')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('.')
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run('If you have any questions concerning this invoice, contact ')
    run.font.size = Pt(10)
    run = p.add_run('{{ contact_name }}')
    run.bold = True
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run('{{ contact_info }}')
    run.font.size = Pt(10)

    doc.save(OUTPUT)
    print(f'Template saved: {OUTPUT}')


if __name__ == '__main__':
    build_template()
