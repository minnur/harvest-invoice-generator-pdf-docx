import csv
import json
import sys
import os
import re
import glob
import subprocess
from datetime import datetime
from docxtpl import DocxTemplate, RichText
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

if len(sys.argv) < 2:
    print('Usage: python3 generate_invoice.py <harvest_csv_or_directory> [config.json]')
    sys.exit(1)

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
CONFIG_PATH = sys.argv[2] if len(sys.argv) > 2 else os.path.join(SCRIPT_DIR, 'config.json')
print(f'Config: {CONFIG_PATH}')
with open(CONFIG_PATH) as f:
    cfg = json.load(f)

LOCAL_CONFIG_PATH = os.path.join(os.path.dirname(CONFIG_PATH), 'local.config.json')
if os.path.exists(LOCAL_CONFIG_PATH):
    print(f'Local config: {LOCAL_CONFIG_PATH}')
    with open(LOCAL_CONFIG_PATH) as f:
        local_cfg = json.load(f)
    for key, val in local_cfg.items():
        if isinstance(val, dict) and isinstance(cfg.get(key), dict):
            cfg[key].update(val)
        else:
            cfg[key] = val

TEMPLATE_PATH = cfg.get('template', os.path.join(SCRIPT_DIR, 'templates', 'invoice-default.docx'))


def parse_csv_dates(filename):
    """Extract start/end dates from Harvest CSV filename.
    Pattern: harvest_time_report_from{YYYY-MM-DD}to{YYYY-MM-DD}.csv
    Returns (start_date, end_date) as datetime objects.
    """
    basename = os.path.basename(filename)
    m = re.match(r'harvest_time_report_from(\d{4}-\d{2}-\d{2})to(\d{4}-\d{2}-\d{2})\.csv', basename)
    if not m:
        return None, None
    start = datetime.strptime(m.group(1), '%Y-%m-%d')
    end = datetime.strptime(m.group(2), '%Y-%m-%d')
    return start, end


def shade_cell(cell, color='F2F2F2'):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    shading.append(shading_elm)


ITEMIZED_MARKER = '___ITEMIZED_TABLE___'


def build_itemized_table(doc, invoice_rows, total_hours, total_amount, rate):
    """Build the itemized table directly in the document, replacing the marker paragraph."""
    # Find the marker paragraph and replace it with the table
    body = doc.element.body
    marker_p = None
    for p in body.findall(qn('w:p')):
        if ITEMIZED_MARKER in ''.join(p.itertext()):
            marker_p = p
            break

    if marker_p is None:
        raise RuntimeError('Could not find itemized table marker in rendered template')

    # Create table after the marker, then remove the marker
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    table.autofit = False

    col_widths = [Inches(1.0), Inches(3.2), Inches(0.8), Inches(0.9), Inches(1.0)]

    # Header row
    header_cells = table.rows[0].cells
    for i, h in enumerate(['DATE', 'DESCRIPTION', 'RATE', 'HOURS', 'AMOUNT']):
        header_cells[i].width = col_widths[i]
        p = header_cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        shade_cell(header_cells[i], 'E8E8E8')

    header_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r in invoice_rows:
        row = table.add_row()
        for i in range(5):
            row.cells[i].width = col_widths[i]

        # Date
        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(r['date'])
        run.font.size = Pt(9)

        # Description (mixed formatting)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        if r['project_code']:
            run = p.add_run(f"[{r['project_code']}]")
            run.bold = True
            run.font.size = Pt(9)
            run = p.add_run(f" {r['notes']}")
            run.font.size = Pt(9)
        else:
            run = p.add_run(r['notes'])
            run.font.size = Pt(9)
        if r['ref_url']:
            p.add_run('\n')
            ref_run = p.add_run(f"Reference: {r['ref_url']}")
            ref_run.font.size = Pt(7)
            ref_run.font.color.rgb = RGBColor(120, 120, 120)

        # Rate
        p = row.cells[2].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(r['rate'])
        run.font.size = Pt(9)

        # Hours
        p = row.cells[3].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{r['hours']:.2f}")
        run.font.size = Pt(9)

        # Amount
        p = row.cells[4].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"${r['amount']:.2f}")
        run.font.size = Pt(9)

    # --- Summary rows ---

    # Subtotal
    row = table.add_row()
    row.cells[0].merge(row.cells[2])
    p = row.cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('Subtotal')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    p = row.cells[3].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'{total_hours:.2f}')
    run.bold = True
    run.font.size = Pt(9)

    p = row.cells[4].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'${total_amount:,.2f}')
    run.bold = True
    run.font.size = Pt(9)

    # Tax Rate, Sales Tax, Other
    for label in ['Tax Rate', 'Sales Tax', 'Other']:
        row = table.add_row()
        row.cells[0].merge(row.cells[3])
        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(label)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

        p = row.cells[4].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run('-')
        run.font.size = Pt(9)

    # Total
    row = table.add_row()
    row.cells[0].merge(row.cells[3])
    p = row.cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('Total')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    p = row.cells[4].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'${total_amount:,.2f}')
    run.bold = True
    run.font.size = Pt(9)

    # Move table to where the marker was, then remove marker
    body.insert(list(body).index(marker_p), table._tbl)
    body.remove(marker_p)


def generate_invoice(csv_path, cfg, invoice_number, invoice_date, output_dir):
    RATE = cfg['rate']
    COL = cfg['csv_columns']
    OUTPUT_PREFIX = cfg['sender']['name'].replace(' ', '-') + '-invoice'
    OUTPUT_DOCX = os.path.join(output_dir, f'{OUTPUT_PREFIX}-{invoice_number}.docx')
    OUTPUT_PDF = os.path.join(output_dir, f'{OUTPUT_PREFIX}-{invoice_number}.pdf')

    if os.path.exists(OUTPUT_DOCX) and os.path.exists(OUTPUT_PDF):
        print(f'DOCX already exists: {OUTPUT_DOCX}')
        print(f'PDF already exists: {OUTPUT_PDF}')
        print('Skipping generation. Delete existing files to regenerate.')
        return

    # --- Step 1: Read Harvest CSV and build invoice rows ---
    invoice_rows = []
    with open(csv_path, 'r') as f:
        reader = csv.reader(f)
        header = next(reader)
        for row in reader:
            date = row[COL['date']]
            project_code = row[COL['project_code']].strip()
            notes = row[COL['notes']].strip()
            hours = float(row[COL['hours']])
            ref_url = row[COL['ref_url']].strip() if len(row) > COL['ref_url'] else ''
            amount = round(hours * RATE, 2)

            invoice_rows.append({
                'date': date,
                'project_code': project_code,
                'notes': notes,
                'ref_url': ref_url,
                'rate': f'${RATE:.2f}',
                'hours': hours,
                'amount': amount,
            })

    total_hours = sum(r['hours'] for r in invoice_rows)
    total_amount = sum(r['amount'] for r in invoice_rows)

    print(f'{len(invoice_rows)} line items, {total_hours:.2f} hours, ${total_amount:,.2f}')

    # --- Step 2: Build RichText objects for header ---

    # Left header: sender info
    left_header = RichText()
    left_header.add(cfg['sender']['name'], bold=True, size=22)
    for line in cfg['sender']['address']:
        left_header.add('\n')
        left_header.add(line, size=20)
    left_header.add('\n')
    left_header.add(f"Email: {cfg['sender']['email']}", size=20)

    # Right header: invoice details + bill to
    right_header = RichText()
    right_header.add('DATE: ', color='808080', size=20)
    right_header.add(invoice_date, bold=True, size=20)
    right_header.add('\n')
    right_header.add('INVOICE#: ', color='808080', size=20)
    right_header.add(invoice_number, bold=True, size=20)
    right_header.add('\n')
    right_header.add('FOR: ', color='808080', size=20)
    right_header.add(cfg['invoice']['for'], size=20)
    right_header.add('\n')
    right_header.add('BILL TO:', color='808080', size=20)
    for line in cfg['bill_to']['company']:
        right_header.add('\n')
        right_header.add(line, bold=True, size=20)
    for line in cfg['bill_to']['address']:
        right_header.add('\n')
        right_header.add(line, size=20)

    # --- Step 3: Render template, then insert itemized table ---
    tpl = DocxTemplate(TEMPLATE_PATH)

    context = {
        'left_header': left_header,
        'right_header': right_header,
        'itemized_table': ITEMIZED_MARKER,
        'payable_to': cfg['footer']['payable_to'],
        'contact_name': cfg['footer']['contact_name'],
        'contact_info': cfg['footer']['contact_info'],
    }

    tpl.render(context)
    doc = tpl.docx

    # Build table directly in the document, replacing the marker paragraph
    build_itemized_table(doc, invoice_rows, total_hours, total_amount, RATE)

    doc.save(OUTPUT_DOCX)
    print(f'DOCX saved: {OUTPUT_DOCX}')

    # Convert DOCX to PDF via Pages
    applescript = f'''
tell application "Pages"
    set doc to open POSIX file "{OUTPUT_DOCX}"
    delay 2
    export doc to POSIX file "{OUTPUT_PDF}" as PDF
    close doc saving no
end tell
'''
    subprocess.run(['osascript', '-e', applescript], check=True)
    subprocess.run(['SetFile', '-a', 'e', OUTPUT_PDF], check=True)
    print(f'PDF saved: {OUTPUT_PDF}')


# --- Main: batch vs single-file mode ---
input_path = sys.argv[1]
if not os.path.isabs(input_path):
    input_path = os.path.join(SCRIPT_DIR, input_path)

if os.path.isdir(input_path):
    # Batch mode: process all matching CSVs in the directory
    csv_files = glob.glob(os.path.join(input_path, 'harvest_time_report_from*to*.csv'))
    if not csv_files:
        print(f'No harvest CSV files found in {input_path}')
        sys.exit(1)

    # Parse dates and sort by start date
    dated_files = []
    for csv_file in csv_files:
        start, end = parse_csv_dates(csv_file)
        if start is None:
            print(f'Warning: skipping {os.path.basename(csv_file)} (could not parse dates)')
            continue
        dated_files.append((start, end, csv_file))

    dated_files.sort(key=lambda x: x[0])

    # Extract year prefix from config invoice number (e.g. "2026-1" -> "2026")
    year_prefix = cfg['invoice']['number'].rsplit('-', 1)[0]

    output_dir = cfg.get('output_dir', '') or input_path

    print(f'Batch mode: {len(dated_files)} CSV files, year prefix: {year_prefix}')
    for seq, (start, end, csv_file) in enumerate(dated_files, start=1):
        invoice_number = f'{year_prefix}-{seq}'
        invoice_date = end.strftime('%B %-d, %Y')
        print(f'\n--- Invoice {invoice_number} from {os.path.basename(csv_file)} ---')
        generate_invoice(csv_file, cfg, invoice_number, invoice_date, output_dir)
else:
    # Single-file mode: original behavior
    generate_invoice(
        input_path,
        cfg,
        cfg['invoice']['number'],
        cfg['invoice']['date'],
        cfg.get('output_dir', '') or os.path.dirname(input_path),
    )
